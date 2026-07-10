# -*- coding: utf-8 -*-
"""
阳光优采平台运营推广月报数据处理脚本
从阳光优采月报 docx 文件中提取指定表格，输出到 Excel 文件。
"""

import os
from datetime import datetime
from docx import Document
import pandas as pd


def extract_tables_from_docx(docx_path):
    """
    从 docx 文件中提取指定表格，返回 {sheet_name: DataFrame} 的字典。
    提取：第1个表(数据总览)、第3个表(分公司表)
    """
    doc = Document(docx_path)
    total = len(doc.tables)
    print(f"文档中共有 {total} 个表格")

    target_tables = [
        (0, "数据总览"),
        (1, "分公司表"),
    ]

    result = {}
    for idx, sheet_name in target_tables:
        if idx < 0 or idx >= total:
            print(f"警告：索引 {idx} 超出范围（共 {total} 个表），跳过 {sheet_name}")
            continue

        table = doc.tables[idx]
        rows = len(table.rows)
        cols = len(table.columns)
        print(f"提取 Table {idx + 1} → {sheet_name}（{rows} 行 x {cols} 列）")

        # 将表格转为二维列表
        data = []
        for row in table.rows:
            data.append([cell.text.strip() for cell in row.cells])

        # 第一行作为表头
        df = pd.DataFrame(data[1:], columns=data[0])
        result[sheet_name] = df

    return result


def save_to_excel(tables_dict, output_path):
    """
    将多个 DataFrame 写入同一个 Excel 文件的不同 sheet。
    """
    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    with pd.ExcelWriter(output_path, engine="openpyxl") as writer:
        for sheet_name, df in tables_dict.items():
            df.to_excel(writer, sheet_name=sheet_name, index=False)
            print(f"已写入 sheet: {sheet_name}")

    print(f"\n输出文件: {output_path}")


def main():
    # 路径配置
    now = datetime.now()
    year = now.year if now.month > 1 else now.year - 1
    month = now.month - 1 if now.month > 1 else 12
    year_month = f"{year}{month:02d}"

    # 动态计算项目根目录
    script_dir = os.path.dirname(os.path.abspath(__file__))
    project_dir = os.path.dirname(script_dir)
    base_dir = os.path.join(project_dir, "Data", year_month)
    source_dir = os.path.join(base_dir, "source_data")
    process_dir = os.path.join(base_dir, "process_data")

    docx_filename = "阳光优采平台运营推广月报.docx"
    output_filename = "yangcai_data.xlsx"

    docx_path = os.path.join(source_dir, docx_filename)
    output_path = os.path.join(process_dir, output_filename)

    if not os.path.exists(docx_path):
        print(f"错误：文件不存在 - {docx_path}")
        return

    print(f"读取文件: {docx_path}\n")
    tables = extract_tables_from_docx(docx_path)
    save_to_excel(tables, output_path)
    print("\n处理完成！")


if __name__ == "__main__":
    main()
