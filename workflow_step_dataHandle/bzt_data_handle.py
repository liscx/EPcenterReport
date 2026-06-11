# -*- coding: utf-8 -*-
"""
标证通月报数据处理脚本
从【标证通】月报 docx 文件中提取指定表格，输出到 Excel 文件。
"""

import os
from datetime import datetime
from docx import Document
import pandas as pd


def extract_tables_from_docx(docx_path):
    """
    从 docx 文件中提取指定表格，返回 {sheet_name: DataFrame} 的字典。
    提取：第1个表(数据总览)、第3个表(涨幅红榜)、第4个表(涨幅绿表)、倒数第3个表(标讯销售分析表)
    """
    doc = Document(docx_path)
    total = len(doc.tables)
    print(f"文档中共有 {total} 个表格")

    # 定义要提取的表格：(索引, sheet名称)
    # 索引从0开始；倒数第三个 = total - 3
    target_tables = [
        (0, "数据总览"),
        (1, "分公司详情"),
        (2, "涨幅红榜"),
        (3, "涨幅绿表"),
        (total - 3, "标讯销售分析表"),
    ]

    result = {}
    for idx, sheet_name in target_tables:
        if idx < 0 or idx >= total:
            print(f"警告：索引 {idx} 超出范围（共 {total} 个表），跳过 {sheet_name}")
            continue

        table = doc.tables[idx]
        rows = len(table.rows)
        cols = len(table.columns)
        print(f"提取 Table {idx + 1} → {sheet_name}（{rows} 行 x {cols} 列）")

        # 将表格转为二维列表
        data = []
        for row in table.rows:
            data.append([cell.text.strip() for cell in row.cells])

        # 第一行作为表头
        df = pd.DataFrame(data[1:], columns=data[0])
        result[sheet_name] = df

    return result


def save_to_excel(tables_dict, output_path):
    """
    将多个 DataFrame 写入同一个 Excel 文件的不同 sheet。
    """
    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    with pd.ExcelWriter(output_path, engine="openpyxl") as writer:
        for sheet_name, df in tables_dict.items():
            df.to_excel(writer, sheet_name=sheet_name, index=False)
            print(f"已写入 sheet: {sheet_name}")

    print(f"\n输出文件: {output_path}")


def main():
    # 路径配置
    now = datetime.now()
    year_month = f"{now.year}{now.month - 1:02d}" if now.month > 1 else f"{now.year - 1}12"
    timestamp = now.strftime("%Y%m%d")
    base_dir = rf"D:\AutoWorkSkill\normalSkills\centerReport\Data\{year_month}"
    source_dir = os.path.join(base_dir, "source_data")
    process_dir = os.path.join(base_dir, "process_data")

    docx_filename = "【标证通】月报.docx"
    output_filename = "bzt_data.xlsx"

    docx_path = os.path.join(source_dir, docx_filename)
    output_path = os.path.join(process_dir, output_filename)

    if not os.path.exists(docx_path):
        print(f"错误：文件不存在 - {docx_path}")
        return

    print(f"读取文件: {docx_path}\n")
    tables = extract_tables_from_docx(docx_path)
    save_to_excel(tables, output_path)
    print("\n处理完成！")


if __name__ == "__main__":
    main()
