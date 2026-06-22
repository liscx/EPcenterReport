# -*- coding: utf-8 -*-
"""
paragraph_export — 从 docx 文件中提取指定段落，保存为 JSON

提取内容：
  - ejy1：e交易月报 段落[4]（全年总收益概况）
  - ejy2：e交易月报 段落[5]（营运/项目费BP完成情况）
  - ejy3：e交易月报 段落[6]（五月产品收益详情）
  - bzt1：标证通月报 段落[3]（产品营收总览）
  - bzt2：标证通月报 段落[7]（新点标证通收益分析）

输出：Data/{year}{month}/res_data/report_paragraphs.json
"""
import os
import sys
import json
from docx import Document

# 路径配置
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, os.path.join(BASE_DIR, 'workflow_step_data_scripts'))
from utils import get_year, get_month

_year = get_year()
_month = get_month()
SOURCE_DIR = os.path.join(BASE_DIR, 'Data', f'{_year}{_month:02d}', 'source_data')
OUTPUT_JSON = os.path.join(BASE_DIR, 'Data', f'{_year}{_month:02d}', 'res_data', 'report_paragraphs.json')

# 各 docx 文件及要提取的段落 {文件名: {段落索引: 键名}}
DOCX_MAP = {
    'e交易收益月报.docx': {4: 'ejy1', 5: 'ejy2', 6: 'ejy3'},
    '【标证通】月报.docx': {3: 'bzt1', 7: 'bzt2', 42: 'bzt3'},
}


def extract_paragraphs():
    """从各 docx 中提取指定段落，合并返回 {键: 段落文本} 字典。"""
    result = {}
    for filename, para_map in DOCX_MAP.items():
        filepath = os.path.join(SOURCE_DIR, filename)
        if not os.path.exists(filepath):
            print(f'文件不存在: {filepath}')
            continue

        doc = Document(filepath)
        for idx, key in para_map.items():
            if idx < len(doc.paragraphs):
                text = doc.paragraphs[idx].text.strip()
                # 英文双引号替换为中文引号，避免破坏 JSON 结构
                text = text.replace('"', '“')
                result[key] = text
            else:
                print(f'{filename} 段落[{idx}]不存在，跳过')
    return result


def process():
    """主处理逻辑。"""
    data = extract_paragraphs()
    if not data:
        print('未提取到任何内容')
        return

    os.makedirs(os.path.dirname(OUTPUT_JSON), exist_ok=True)
    with open(OUTPUT_JSON, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    print(f'段落已保存: {OUTPUT_JSON}')


if __name__ == '__main__':
    process()
