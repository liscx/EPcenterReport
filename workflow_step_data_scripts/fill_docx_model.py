# -*- coding: utf-8 -*-
"""
fill_docx_model — 将 extract_data 月报 Excel 数据填入 docx 模型文件

逻辑：
  1. 读取 model/运营中心营运产品收益月报_model.docx（仅有表头的模板）
  2. 读取 Data/{year}{month}/res_data/extract_data{month}月报.xlsx
  3. 按 sheet 名（表格N）映射到 docx 的第 N-1 个表格
  4. 将 Excel 数据逐行追加为 docx 表格的新行
  5. 输出到 Data/{year}{month}/res_data/运营中心营运产品收益月报.docx

注意事项：
  - docx 表格仅保留一行表头，数据从第二行开始写入
  - 合并单元格的表格（表格8、表格19）按原始合并结构处理
  - 表格2（e交易-分公司收益）同一分公司的运营/项目费行，全年收益/BP完成比例纵向合并
  - Excel 中不存在的 sheet 对应的 docx 表格保持原样（仅表头）
"""
import os
import json
import pandas as pd
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy
from utils import get_month, get_year, BASE_DIR, format_number

# ── 路径配置 ──────────────────────────────────────────────────────────
MODEL_FILE = os.path.join(BASE_DIR, 'model', '运营中心营运产品收益月报_model.docx')

_month = get_month()
_year = get_year()
RES_DATA_DIR = os.path.join(BASE_DIR, 'Data', f'{_year}{_month:02d}', 'res_data')
EXTRACT_FILE = os.path.join(RES_DATA_DIR, f'extract_data{_month}月报.xlsx')
PARAGRAPH_JSON = os.path.join(RES_DATA_DIR, 'report_paragraphs.json')
OUTPUT_DOCX = os.path.join(RES_DATA_DIR, '运营中心营运产品收益月报.docx')


def create_plain_row(table):
    """创建一个无格式的空数据行（无底纹、无加粗等表头样式）。"""
    num_cols = len(table.columns)
    # 从最后一行表头复制列宽结构（避免合并单元格问题）
    source_tr = table.rows[-1]._tr
    new_tr = deepcopy(source_tr)

    # 清除每个单元格的格式和文本
    for tc in new_tr.findall(qn('w:tc')):
        # 移除合并单元格属性（如果有）
        tc_pr = tc.find(qn('w:tcPr'))
        if tc_pr is not None:
            # 移除 gridSpan（横向合并）
            grid_span = tc_pr.find(qn('w:gridSpan'))
            if grid_span is not None:
                tc_pr.remove(grid_span)
            # 移除 vMerge（纵向合并）
            v_merge = tc_pr.find(qn('w:vMerge'))
            if v_merge is not None:
                tc_pr.remove(v_merge)
            # 保留列宽，移除底纹
            shd = tc_pr.find(qn('w:shd'))
            if shd is not None:
                tc_pr.remove(shd)

        # 清除单元格内所有段落和 run，只保留一个空段落
        for p in tc.findall(qn('w:p')):
            # 移除段落级样式
            p_pr = p.find(qn('w:pPr'))
            if p_pr is not None:
                p_style = p_pr.find(qn('w:pStyle'))
                if p_style is not None:
                    p_pr.remove(p_style)
            # 移除所有 run
            for r in p.findall(qn('w:r')):
                p.remove(r)

    return new_tr


def add_data_row(table, row_data, num_cols, color_map=None):
    """
    向 docx 表格追加一行无格式的数据。

    Args:
        color_map: dict {col_idx: 'RRGGBB'}，为指定列的文本设置字体颜色。
    """
    new_tr = create_plain_row(table)
    table._tbl.append(new_tr)

    new_row = table.rows[-1]
    for col_idx in range(min(num_cols, len(new_row.cells))):
        val = row_data[col_idx] if col_idx < len(row_data) else ''
        cell = new_row.cells[col_idx]
        # 数字统一格式化：保留2位小数，整数不留小数点
        if not pd.isna(val) and pd.api.types.is_number(val):
            text = format_number(val)
        else:
            text = str(val) if not pd.isna(val) else ''
        # 清除单元格所有段落，只保留一个空段落
        for p in cell.paragraphs:
            for r in p.runs:
                r.text = ''
        # 设置文本
        if cell.paragraphs:
            # 清除所有现有 run
            for r in cell.paragraphs[0].runs:
                r.text = ''
            # 如果没有 run，添加一个
            if not cell.paragraphs[0].runs:
                run = cell.paragraphs[0].add_run(text)
            else:
                cell.paragraphs[0].runs[0].text = text
            # 设置字体颜色
            if color_map and col_idx in color_map:
                run_obj = cell.paragraphs[0].runs[0]
                rpr = run_obj._r.get_or_add_rPr()
                color_el = OxmlElement('w:color')
                color_el.set(qn('w:val'), color_map[col_idx])
                rpr.append(color_el)


def merge_cells_vertical(table, start_row, count, col_idx):
    """将 table 中从 start_row 开始的连续 count 行的第 col_idx 列合并为一个单元格，并设置居中对齐。"""
    for i in range(count):
        cell = table.cell(start_row + i, col_idx)
        tc = cell._tc
        tc_pr = tc.get_or_add_tcPr()
        if i == 0:
            # 首行：开始合并
            v_merge = OxmlElement('w:vMerge')
            v_merge.set(qn('w:val'), 'restart')
            tc_pr.append(v_merge)
        else:
            # 后续行：继续合并
            v_merge = OxmlElement('w:vMerge')
            tc_pr.append(v_merge)

    # 设置首行单元格的垂直居中对齐
    first_cell = table.cell(start_row, col_idx)
    for p in first_cell.paragraphs:
        p.alignment = 1  # 1 = center（水平居中）
        # 设置垂直居中
        pPr = p._p.get_or_add_pPr()
        vAlign = OxmlElement('w:vAlign')
        vAlign.set(qn('w:val'), 'center')
        pPr.append(vAlign)


def align_columns(df, table):
    """
    直接按顺序返回 DataFrame，不做列名匹配。
    """
    return df


def fill_paragraph_placeholders(doc, json_file):
    """将 docx 段落中的 {{key}} 占位符替换为 JSON 中的文本。"""
    if not os.path.exists(json_file):
        print(f'段落数据文件不存在，跳过占位符替换: {json_file}')
        return

    with open(json_file, 'r', encoding='utf-8') as f:
        data = json.load(f)

    count = 0
    for p in doc.paragraphs:
        for key, value in data.items():
            placeholder = '{{' + key + '}}'
            if placeholder in p.text:
                # 遍历 runs 替换文本（保留格式）
                for run in p.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, value)
                        count += 1
                # 如果占位符跨了多个 runs，用完整段落文本替换
                if placeholder in p.text:
                    full_text = p.text.replace(placeholder, value)
                    # 清空所有 runs，把替换后的文本写入第一个 run
                    if p.runs:
                        for r in p.runs:
                            r.text = ''
                        p.runs[0].text = full_text
                        count += 1

    print(f'占位符替换完成，共替换 {count} 处')


def fill_table_from_excel(table, df, merge_cols=None, key_col=0, color_cols=None):
    """
    将 DataFrame 数据填入 docx 表格。

    Args:
        merge_cols: 需要纵向合并单元格的列索引列表。
                    当连续行的 key_col 列值相同时，这些列的单元格会被合并。
        key_col: 用于判断合并的列索引（默认为首列）。
        color_cols: 需要根据值标色的列索引列表。
                    ▲ 开头标红色，▼ 开头标绿色。
    """
    num_cols = len(table.columns)
    for _, row in df.iterrows():
        row_data = [row.iloc[i] if i < len(row) else '/' for i in range(num_cols)]
        # 根据值动态生成颜色映射
        color_map = {}
        if color_cols:
            for col_idx in color_cols:
                if col_idx < len(row_data):
                    val_str = str(row_data[col_idx])
                    if val_str.startswith('▲'):
                        color_map[col_idx] = 'FF0000'  # 红色
                    elif val_str.startswith('▼'):
                        color_map[col_idx] = '00B050'  # 绿色
        add_data_row(table, row_data, num_cols, color_map=color_map)

    # 纵向合并单元格（按指定列值判断连续相同行）
    if merge_cols:
        header_rows = len(table.rows) - len(df)  # 表头行数
        total_rows = len(table.rows)
        i = header_rows
        while i < total_rows:
            key = table.cell(i, key_col).text.strip()
            count = 1
            while i + count < total_rows and table.cell(i + count, key_col).text.strip() == key:
                count += 1
            if count > 1:
                for col_idx in merge_cols:
                    merge_cells_vertical(table, i, count, col_idx)
            i += count


def fill_table_8(table, df):
    """
    表格8（阳光优采-分公司专区情况）特殊处理：
    docx 模型有 2 行表头（含合并单元格），需要按子表头列顺序写入数据。
    Excel 表格8 的列：序号, 分公司, 本月专区情况(3列), 上月专区情况(3列), 合计情况(3列), 全年收益, BP总额, BP完成率
    """
    num_cols = 14  # 固定14列
    for _, row in df.iterrows():
        row_data = [row.iloc[i] if i < len(row) else '/' for i in range(num_cols)]
        add_data_row(table, row_data, num_cols)


def fill_table_19(table, df, color_cols=None):
    """
    表格19（投标保函-平台销售情况）特殊处理：
    docx 模型只有表头行，需要追加新行。
    """
    num_cols = len(table.columns)
    for _, row in df.iterrows():
        row_data = [row.iloc[i] if i < len(row) else '/' for i in range(num_cols)]
        # 根据值动态生成颜色映射
        color_map = {}
        if color_cols:
            for col_idx in color_cols:
                if col_idx < len(row_data):
                    val_str = str(row_data[col_idx])
                    if val_str.startswith('▲'):
                        color_map[col_idx] = 'FF0000'  # 红色
                    elif val_str.startswith('▼'):
                        color_map[col_idx] = '00B050'  # 绿色
        add_data_row(table, row_data, num_cols, color_map=color_map)


def fill_table_7(table, df, color_cols=None):
    """
    表格7（阳光优采-核心指标）特殊处理：
    docx 模型有表头行（合并2列）和空行，需要填充到空行中。
    第一列相同值会纵向合并居中。
    """
    num_cols = len(table.columns)
    rows = table.rows

    # 从第2行开始填充（跳过表头行）
    start_row = 1
    for idx, (_, row) in enumerate(df.iterrows()):
        row_data = [row.iloc[i] if i < len(row) else '/' for i in range(num_cols)]

        # 根据值动态生成颜色映射
        color_map = {}
        if color_cols:
            for col_idx in color_cols:
                if col_idx < len(row_data):
                    val_str = str(row_data[col_idx])
                    if val_str.startswith('▲'):
                        color_map[col_idx] = 'FF0000'  # 红色
                    elif val_str.startswith('▼'):
                        color_map[col_idx] = '00B050'  # 绿色

        # 如果当前行存在，填充到该行
        if start_row + idx < len(rows):
            current_row = rows[start_row + idx]
            for col_idx, value in enumerate(row_data):
                if col_idx < len(current_row.cells):
                    cell = current_row.cells[col_idx]
                    # 清空单元格内容
                    for p in cell.paragraphs:
                        for r in p.runs:
                            r.text = ''
                    # 设置新值
                    if cell.paragraphs:
                        p = cell.paragraphs[0]
                        if p.runs:
                            run = p.runs[0]
                        else:
                            run = p.add_run()
                        run.text = str(value)
                        # 设置颜色
                        if color_map and col_idx in color_map:
                            from docx.oxml.ns import qn
                            from docx.oxml import OxmlElement
                            rpr = run._r.get_or_add_rPr()
                            color_el = OxmlElement('w:color')
                            color_el.set(qn('w:val'), color_map[col_idx])
                            rpr.append(color_el)
        else:
            # 如果当前行不存在，添加新行
            add_data_row(table, row_data, num_cols, color_map=color_map)

    # 第一列相同值纵向合并居中
    total_rows = len(table.rows)
    i = start_row
    while i < total_rows:
        key = table.cell(i, 0).text.strip()
        count = 1
        while i + count < total_rows and table.cell(i + count, 0).text.strip() == key:
            count += 1
        if count > 1:
            merge_cells_vertical(table, i, count, 0)
        i += count


def process():
    print(f'模板文件: {MODEL_FILE}')
    print(f'数据文件: {EXTRACT_FILE}')
    print(f'输出文件: {OUTPUT_DOCX}')

    if not os.path.exists(MODEL_FILE):
        print(f'错误: 模板文件不存在: {MODEL_FILE}')
        return
    if not os.path.exists(EXTRACT_FILE):
        print(f'错误: 数据文件不存在: {EXTRACT_FILE}')
        return

    # 加载 docx 模板
    doc = Document(MODEL_FILE)

    # 加载 Excel 数据
    xls = pd.ExcelFile(EXTRACT_FILE, engine='openpyxl')
    sheet_names = xls.sheet_names
    print(f'Excel sheets: {sheet_names}')

    # sheet 名 → 表格索引的映射
    # "表格N" → docx.tables[N-1]
    for sheet_name in sheet_names:
        if not sheet_name.startswith('表格'):
            continue

        try:
            table_num = int(sheet_name.replace('表格', ''))
        except ValueError:
            print(f'跳过无法解析的 sheet: {sheet_name}')
            continue

        table_idx = table_num - 1  # 表格1 → index 0

        if table_idx < 0 or table_idx >= len(doc.tables):
            print(f'警告: {sheet_name} 映射到表格索引 {table_idx}，超出范围（共 {len(doc.tables)} 个表格），跳过')
            continue

        df = pd.read_excel(xls, sheet_name=sheet_name)
        if df.empty:
            print(f'{sheet_name}: 数据为空，跳过')
            continue

        table = doc.tables[table_idx]

        # 按 docx 表头对齐列顺序，缺失列填 '/'
        df = align_columns(df, table)
        print(f'{sheet_name} → docx 表格 #{table_idx}（{len(df)} 行数据, {len(table.columns)} 列）')

        # 自动检测环比/同比列索引，用于标色
        color_cols = []
        for i, col in enumerate(df.columns):
            if any(k in str(col) for k in ['环比', '同比']):
                color_cols.append(i)

        # 表格8 和 表格19 有合并单元格，需要特殊处理
        if table_num == 8:
            fill_table_8(table, df)
        elif table_num == 19:
            fill_table_19(table, df, color_cols=color_cols)
        elif table_num == 1:
            # 表格1：产品线列需要合并居中
            fill_table_from_excel(table, df, merge_cols=[0], key_col=0, color_cols=color_cols)
        elif table_num == 2:
            # 表格2：同一分公司的运营/项目费行，合并全年收益、BP完成比例
            fill_table_from_excel(table, df, merge_cols=[7, 9], key_col=1, color_cols=color_cols)
        elif table_num == 7:
            # 表格7：表头合并2列，需要填充到空行中
            fill_table_7(table, df, color_cols=color_cols)
        else:
            fill_table_from_excel(table, df, color_cols=color_cols if color_cols else None)

    # 替换段落占位符（{{ejy1}} 等）
    fill_paragraph_placeholders(doc, PARAGRAPH_JSON)

    # 保存输出文件
    os.makedirs(RES_DATA_DIR, exist_ok=True)
    doc.save(OUTPUT_DOCX)
    print(f'已保存: {OUTPUT_DOCX}')


if __name__ == '__main__':
    process()
